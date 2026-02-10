# file_generator.py
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
import re
import unicodedata

EXPORT_DIR = Path("exports")
EXPORT_DIR.mkdir(exist_ok=True)

CHOICE_PATTERN = re.compile(r"^\s*(\(?[A-D]\)?[.)])\s+(.+)$")  # A)  A.  (A)  A)
QUESTION_PATTERN = re.compile(r"^\s*\d+\.\s+(.+)$", re.IGNORECASE)


def _sanitize_filename(name: str) -> str:
    name = unicodedata.normalize("NFKD", name)
    name = re.sub(r"[^\w\-\.]+", "_", name)
    return name[:60] if name else "quiz"


def generate_docx_from_text(
    quiz_text: str,
    filename_hint: str | None = None,
    include_answers: bool = True,  # NEW
) -> str:
    """
    Convert plain text quiz into a nicely formatted DOCX:
    - Each non-empty line → paragraph
    - Numbered questions (1.) → bold number
    - Choices A–D → indented
    - "ANSWER KEY" → subheader (optional)
    Returns filesystem path to the saved .docx
    """
    doc = Document()

    # Page styling
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Title (optional)
    title = (filename_hint or "Generated Quiz").strip()
    if title:
        p = doc.add_paragraph(title)
        r = p.runs[0]
        r.font.bold = True
        r.font.size = Pt(16)
        p.paragraph_format.space_after = Pt(12)

    lines = (quiz_text or "").splitlines()
    in_answer_key = False

    for raw in lines:
        line = raw.rstrip()

        # If we don't want answers: stop when answer key starts
        if not include_answers and line.strip().lower().startswith("answer key"):
            break

        # Blank line → spacer
        if not line.strip():
            para = doc.add_paragraph("")
            para.paragraph_format.space_after = Pt(6)
            continue

        # Detect ANSWER KEY header (case-insensitive)
        if line.strip().lower().startswith("answer key"):
            if include_answers:
                in_answer_key = True
                p = doc.add_paragraph("ANSWER KEY")
                r = p.runs[0]
                r.font.bold = True
                r.font.size = Pt(12)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
            continue

        # Inside answer key
        if in_answer_key:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(3)
            continue

        # Question? (e.g., "1. ...")
        q_match = QUESTION_PATTERN.match(line)
        if q_match:
            p = doc.add_paragraph()
            q_run_num = p.add_run(line.split(".", 1)[0] + ". ")
            q_run_num.bold = True
            q_run_num.font.size = Pt(11)

            q_text = line.split(".", 1)[1].strip() if "." in line else line
            q_run_text = p.add_run(q_text)
            q_run_text.font.size = Pt(11)

            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            continue

        # Choice? (A) / A. / A)
        c_match = CHOICE_PATTERN.match(line)
        if c_match:
            p = doc.add_paragraph(f"{c_match.group(1)} {c_match.group(2)}")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(2)
            continue

        # Default
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(3)

    # Save
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = _sanitize_filename(filename_hint or "quiz_export")
    out_path = EXPORT_DIR / f"{base}_{stamp}.docx"
    doc.save(out_path)
    return str(out_path)
